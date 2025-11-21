import functions_framework
import json
import os
import logging
from flask import request
from google.cloud import storage
import google.generativeai as genai
import elasticsearch
from PyPDF2 import PdfReader
from docx import Document
import hashlib
import tempfile
from datetime import datetime
import base64

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Configuration ---
PROJECT_ID = os.environ.get('GCP_PROJECT_ID', 'college-counselling-478115')
REGION = os.environ.get('REGION', 'us-east1')
BUCKET_NAME = os.environ.get('GCS_BUCKET_NAME', 'college-counselling-knowledge-base')
DATA_STORE = os.environ.get('DATA_STORE', 'college_admissions_kb')

# Elasticsearch Configuration
ES_CLOUD_ID = os.environ.get('ES_CLOUD_ID')
ES_API_KEY = os.environ.get('ES_API_KEY')
ES_INDEX_NAME = os.environ.get('ES_INDEX_NAME', 'university_documents')

# Gemini Configuration
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')

# --- CORS Headers ---
def add_cors_headers(response, status_code=200):
    """Add CORS headers to response and return proper format."""
    if isinstance(response, dict):
        response = (json.dumps(response), status_code, {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'GET, POST, PUT, DELETE, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type, Authorization, X-User-Email',
            'Access-Control-Max-Age': '3600'
        })
    elif isinstance(response, tuple) and len(response) >= 2:
        # Add CORS headers to existing tuple response
        if len(response) == 2:
            data, status = response
            headers = {}
        else:
            data, status, headers = response
        
        # Ensure headers is a dict
        if not headers:
            headers = {}
        
        # Add CORS headers
        headers.update({
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'GET, POST, PUT, DELETE, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type, Authorization, X-User-Email',
            'Access-Control-Max-Age': '3600'
        })
        
        response = (data, status, headers)
    
    return response

# --- Client Initialization ---
def get_storage_client():
    """Initialize Google Cloud Storage client."""
    return storage.Client(project=PROJECT_ID)

def get_elasticsearch_client():
    """Initialize Elasticsearch client."""
    if not ES_CLOUD_ID or not ES_API_KEY:
        raise ValueError("Elasticsearch credentials not configured")
    
    return elasticsearch.Elasticsearch(
        cloud_id=ES_CLOUD_ID,
        api_key=ES_API_KEY
    )

def get_gemini_client():
    """Initialize Gemini client."""
    if not GEMINI_API_KEY:
        raise ValueError("Gemini API key not configured")
    
    genai.configure(api_key=GEMINI_API_KEY)
    return genai.GenerativeModel('gemini-2.5-flash')

def get_embedding_client():
    """Initialize Gemini embedding client."""
    if not GEMINI_API_KEY:
        raise ValueError("Gemini API key not configured")
    
    genai.configure(api_key=GEMINI_API_KEY)
    return genai.embed_content(model='models/text-embedding-004')

# --- Elasticsearch Index Management ---
def delete_elasticsearch_index():
    """Delete Elasticsearch index (use with caution!)."""
    try:
        es_client = get_elasticsearch_client()
        if es_client.indices.exists(index=ES_INDEX_NAME):
            es_client.indices.delete(index=ES_INDEX_NAME)
            logger.info(f"Deleted index {ES_INDEX_NAME}")
            return True
        else:
            logger.info(f"Index {ES_INDEX_NAME} does not exist")
            return False
    except Exception as e:
        logger.error(f"Error deleting index: {str(e)}")
        return False

def create_elasticsearch_index(force_recreate=False):
    """Create Elasticsearch index with proper mapping."""
    es_client = get_elasticsearch_client()
    
    if es_client.indices.exists(index=ES_INDEX_NAME):
        if force_recreate:
            logger.info(f"Force recreating index {ES_INDEX_NAME}")
            delete_elasticsearch_index()
        else:
            logger.info(f"Index {ES_INDEX_NAME} already exists")
            return True
    
    mapping = {
        "mappings": {
            "properties": {
                "document_id": {"type": "keyword"},
                "user_id": {"type": "keyword"},
                "filename": {"type": "text", "analyzer": "standard"},
                "content": {
                    "type": "text",
                    "analyzer": "standard"
                },
                "university_name": {"type": "text", "analyzer": "standard"},
                "metadata": {
                    "properties": {
                        "university_identity": {
                            "properties": {
                                "name": {"type": "text", "analyzer": "standard"},
                                "type": {"type": "keyword"},
                                "location": {"type": "text"},
                                "ranking": {"type": "keyword"}
                            }
                        },
                        "academic_structure": {
                            "properties": {
                                "majors_inventory": {"type": "text"},
                                "programs_offered": {"type": "text"},
                                "schools_colleges": {"type": "text"}
                            }
                        },
                        "admissions_statistics": {
                            "properties": {
                                "acceptance_rate": {"type": "float"},
                                "total_applications": {"type": "integer"},
                                "enrollment_size": {"type": "integer"}
                            }
                        }
                    }
                },
                "chunks": {
                    "type": "nested",
                    "properties": {
                        "text": {"type": "text", "analyzer": "standard"},
                        "embedding": {
                            "type": "dense_vector",
                            "dims": 768,
                            "index": True,
                            "similarity": "cosine"
                        },
                        "chunk_index": {"type": "integer"}
                    }
                },
                "num_chunks": {"type": "integer"},
                "status": {"type": "keyword"},
                "processing_stage": {"type": "keyword"},
                "indexed_at": {"type": "date"},
                "file_size": {"type": "long"},
                "file_type": {"type": "keyword"}
            }
        },
        "settings": {
            "number_of_shards": 1,
            "number_of_replicas": 0,
            "refresh_interval": "1s"
        }
    }
    
    try:
        es_client.indices.create(index=ES_INDEX_NAME, body=mapping)
        logger.info(f"Created index {ES_INDEX_NAME}")
        return True
    except Exception as e:
        logger.error(f"Error creating index: {str(e)}")
        return False

# --- Document AI Layout Parser ---
def get_document_ai_client():
    """Initialize Document AI client."""
    from google.cloud import documentai
    return documentai.DocumentProcessorServiceClient()

def get_text_from_layout(layout, text):
    """Extract text from a layout element using text anchors."""
    if not layout.text_anchor:
        return ""
    
    # Get text segments from the layout
    response_text = ""
    for segment in layout.text_anchor.text_segments:
        start_index = int(segment.start_index) if segment.start_index else 0
        end_index = int(segment.end_index) if segment.end_index else len(text)
        response_text += text[start_index:end_index]
    
    return response_text

def process_layout_parser_response(document_layout):
    """Process Document AI Layout Parser response with documentLayout structure."""
    chunks = []
    full_text_parts = []
    
    def extract_blocks(blocks, depth=0):
        """Recursively extract text blocks from nested structure."""
        for block in blocks:
            if hasattr(block, 'text_block') and block.text_block:
                text_block = block.text_block
                text = text_block.text if hasattr(text_block, 'text') else ""
                block_type = text_block.type if hasattr(text_block, 'type') else "paragraph"
                
                if text and text.strip():
                    # Add to full text
                    full_text_parts.append(text)
                    
                    # Create chunk for meaningful blocks (skip bullets and very short text)
                    if len(text.strip()) > 2 and text.strip() not in ['•', '-', '*']:
                        page_start = 1
                        if hasattr(block, 'page_span') and block.page_span:
                            page_start = block.page_span.page_start if hasattr(block.page_span, 'page_start') else 1
                        
                        chunks.append({
                            'text': text.strip(),
                            'type': block_type,
                            'page': page_start,
                            'depth': depth
                        })
                
                # Process nested blocks
                if hasattr(text_block, 'blocks') and text_block.blocks:
                    extract_blocks(text_block.blocks, depth + 1)
    
    # Extract all blocks
    if hasattr(document_layout, 'blocks') and document_layout.blocks:
        extract_blocks(document_layout.blocks)
    
    full_text = '\n\n'.join(full_text_parts)
    
    logger.info(f"Layout Parser extracted {len(chunks)} text blocks, total text length: {len(full_text)}")
    
    return {
        'chunks': chunks,
        'full_text': full_text
    }

def process_document_with_layout(file_path, file_type):
    """Process document using Document AI layout parser for intelligent chunking."""
    try:
        # Use Document AI for PDF files
        if file_type.lower() == 'pdf':
            logger.info("Using Document AI layout parser for PDF")
            from google.cloud import documentai_v1 as documentai
            
            # Initialize Document AI client
            client = documentai.DocumentProcessorServiceClient()
            
            # Document AI processor configuration
            project_id = PROJECT_ID
            location = "us"
            processor_id = "448ced96febb71f7"  # Layout parser processor
            
            # Build the processor name
            name = f"projects/{project_id}/locations/{location}/processors/{processor_id}"
            
            # Read the file
            with open(file_path, 'rb') as file:
                file_content = file.read()
            
            logger.info(f"Processing PDF with Document AI Layout Parser, size: {len(file_content)} bytes")
            logger.info(f"Using processor: {name}")
            
            # Create the document
            raw_document = documentai.RawDocument(
                content=file_content,
                mime_type="application/pdf"
            )
            
            # Process the document with layout parser
            request = documentai.ProcessRequest(
                name=name,
                raw_document=raw_document,
                skip_human_review=True
            )
            
            try:
                result = client.process_document(request=request)
                document = result.document
                
                # Document AI Layout Parser returns standard document response with enhanced layout info
                logger.info(f"Document AI Layout Parser response - Pages: {len(document.pages) if document.pages else 0}")
                logger.info(f"Document has text: {len(document.text) if document.text else 0} chars")
                
                # Use Layout Parser enhanced text extraction
                if document.text and len(document.text.strip()) > 0:
                    logger.info("Using Document AI Layout Parser extracted text")
                    return document.text.strip()
                
                # Fallback to page-based extraction
                logger.info("No full text, extracting from pages")
                
                # Check if we got valid results
                if not document.pages or len(document.pages) == 0:
                    logger.warning("Document AI returned 0 pages, falling back to simple extraction")
                    return extract_text_fallback(file_path, file_type)
                
                # Extract text chunks based on layout blocks
                chunks = []
                full_text = document.text
                
                # Get layout blocks from all pages
                for page_num, page in enumerate(document.pages):
                    for block_num, block in enumerate(page.blocks):
                        # Extract text for this block using layout information
                        block_text = get_text_from_layout(block.layout, document.text)
                        if block_text and block_text.strip():
                            chunks.append({
                                'text': block_text.strip(),
                                'type': 'block',
                                'page': page_num + 1,
                                'block': block_num
                            })
                
                logger.info(f"Extracted {len(chunks)} layout blocks from Document AI")
                
                # If no chunks extracted, use full text as fallback
                if not chunks:
                    logger.warning("No layout blocks extracted, using full text")
                    if full_text and full_text.strip():
                        return full_text.strip()
                    else:
                        return extract_text_fallback(file_path, file_type)
                
                return {
                    'chunks': chunks,
                    'full_text': full_text
                }
            except Exception as doc_ai_error:
                logger.error(f"Document AI processing error: {str(doc_ai_error)}")
                logger.info("Falling back to simple text extraction")
                return extract_text_fallback(file_path, file_type)
        else:
            # For non-PDF files, use fallback
            logger.info(f"Using fallback extraction for {file_type}")
            return extract_text_fallback(file_path, file_type)
        
        # Uncomment when you have Document AI processor configured:
        """
        from google.cloud import documentai
        
        client = get_document_ai_client()
        name = client.processor_path(project_id, location, processor_id)
        
        # Read the file
        with open(file_path, 'rb') as file:
            content = file.read()
        
        # Create the request
        raw_document = documentai.RawDocument(
            content=content,
            mime_type="application/pdf" if file_type.lower() == 'pdf' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        request = documentai.ProcessRequest(name=name, raw_document=raw_document)
        result = client.process_document(request=request)
        
        # Extract text chunks based on layout
        chunks = []
        document = result.document
        
        # Get text blocks with layout information
        for page in document.pages:
            for block in page.blocks:
                block_text = layout_to_text(block.layout, document.text)
                if block_text.strip():
                    chunks.append({
                        'text': block_text.strip(),
                        'type': 'block',
                        'page': page.page_number
                    })
        
        return {
            'chunks': chunks,
            'full_text': document.text
        }
        """
        
    except Exception as e:
        logger.error(f"Error with Document AI layout parser: {str(e)}")
        # Fallback to simple text extraction
        return extract_text_fallback(file_path, file_type)

def extract_text_fallback(file_path, file_type):
    """Fallback text extraction for PDF, DOCX, or TXT files."""
    try:
        if file_type.lower() == 'pdf':
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        elif file_type.lower() == 'docx':
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        elif file_type.lower() in ['txt', 'text']:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            return None
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        return None

def generate_content_embeddings(document_data):
    """Generate vector embeddings for document chunks using Gemini."""
    try:
        import google.generativeai as genai
        
        # Handle both string content and Document AI structured data
        if isinstance(document_data, str):
            logger.info(f"Chunking text document of length {len(document_data)} characters")
            # Simple string content - chunk it intelligently
            chunks = chunk_text_intelligently(document_data)
            logger.info(f"Created {len(chunks)} chunks from text")
        elif isinstance(document_data, dict) and 'chunks' in document_data:
            # Document AI structured data
            chunks = document_data['chunks']
            logger.info(f"Using {len(chunks)} chunks from Document AI")
        else:
            # Fallback: treat as string
            logger.info(f"Fallback chunking for data type: {type(document_data)}")
            chunks = chunk_text_intelligently(str(document_data))
        
        embeddings = []
        max_chunk_size = 8000  # characters, conservative for Gemini limits
        
        for i, chunk in enumerate(chunks):
            if isinstance(chunk, dict):
                chunk_text = chunk['text']
            else:
                chunk_text = chunk
            
            # Further chunk if still too large
            if len(chunk_text) > max_chunk_size:
                chunk_text = chunk_text[:max_chunk_size]
                logger.info(f"Chunk {i} further truncated from {len(chunk_text)} to {max_chunk_size} characters")
            
            # Generate embedding for this chunk
            result = genai.embed_content(
                model='models/text-embedding-004',
                content=chunk_text,
                task_type="retrieval_document"
            )
            
            embeddings.append({
                'text': chunk_text,
                'embedding': result['embedding'],
                'chunk_index': i
            })
            
            logger.info(f"Generated embedding for chunk {i} ({len(chunk_text)} chars)")
        
        logger.info(f"Generated {len(embeddings)} total embeddings for document")
        return embeddings
        
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        return None

def chunk_text_intelligently(text, max_chunk_size=8000):
    """Intelligently chunk text by paragraphs and sentences."""
    chunks = []
    
    # If text is small enough, return as single chunk
    if len(text) <= max_chunk_size:
        return [text]
    
    # Try splitting by double newlines (paragraphs)
    paragraphs = text.split('\n\n')
    
    # If no paragraph breaks, try single newlines
    if len(paragraphs) == 1:
        paragraphs = text.split('\n')
    
    # If still no breaks, force chunk by size
    if len(paragraphs) == 1:
        logger.info(f"No natural breaks found, force chunking {len(text)} chars into {max_chunk_size} char chunks")
        for i in range(0, len(text), max_chunk_size):
            chunks.append(text[i:i + max_chunk_size])
        return chunks
    
    current_chunk = ""
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        # If paragraph itself is too long, split by sentences
        if len(paragraph) > max_chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            
            # Split long paragraph by sentences
            sentences = paragraph.split('. ')
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                # If sentence itself is too long, force chunk it
                if len(sentence) > max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        current_chunk = ""
                    # Force chunk the long sentence
                    for i in range(0, len(sentence), max_chunk_size):
                        chunks.append(sentence[i:i + max_chunk_size])
                elif len(current_chunk) + len(sentence) + 2 <= max_chunk_size:
                    current_chunk += sentence + '. '
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + '. '
        else:
            # Normal paragraph
            separator = '\n\n' if '\n\n' in text else '\n'
            if len(current_chunk) + len(paragraph) + len(separator) <= max_chunk_size:
                current_chunk += paragraph + separator
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + separator
    
    # Add the last chunk
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    logger.info(f"Chunked {len(text)} chars into {len(chunks)} chunks")
    return chunks

def generate_document_metadata(content, filename):
    """Generate structured metadata using Gemini LLM."""
    try:
        gemini_model = get_gemini_client()
        
        prompt = f"""
        Analyze the following university document content and extract structured information:
        
        Filename: {filename}
        Content: {content[:2000]}...
        
        Please extract and categorize the information into this JSON structure:
        {{
            "university_identity": {{
                "name": "University name if mentioned",
                "type": "Public/Private/Unknown",
                "location": "City, State if mentioned",
                "ranking": "National/Regional/Unknown if mentioned"
            }},
            "academic_structure": {{
                "majors_inventory": ["List of majors mentioned"],
                "programs_offered": ["Key programs highlighted"],
                "schools_colleges": ["Academic divisions mentioned"]
            }},
            "admissions_statistics": {{
                "acceptance_rate": "percentage if mentioned",
                "total_applications": "number if mentioned",
                "enrollment_size": "number if mentioned"
            }}
        }}
        
        Return only valid JSON that can be parsed.
        """
        
        response = gemini_model.generate_content(prompt)
        
        # Try to parse the response as JSON
        try:
            import json
            metadata = json.loads(response.text)
            return metadata
        except json.JSONDecodeError:
            # Fallback if response is not valid JSON
            return {
                "university_identity": {"name": "Unknown", "type": "Unknown", "location": "Unknown", "ranking": "Unknown"},
                "academic_structure": {"majors_inventory": [], "programs_offered": [], "schools_colleges": []},
                "admissions_statistics": {"acceptance_rate": None, "total_applications": None, "enrollment_size": None}
            }
            
    except Exception as e:
        logger.error(f"Error generating metadata: {str(e)}")
        return {
            "university_identity": {"name": "Unknown", "type": "Unknown", "location": "Unknown", "ranking": "Unknown"},
            "academic_structure": {"majors_inventory": [], "programs_offered": [], "schools_colleges": []},
            "admissions_statistics": {"acceptance_rate": None, "total_applications": None, "enrollment_size": None}
        }

# --- Document Operations ---
def index_document(user_id, filename, file_content=None, file_path=None):
    """Index document into Elasticsearch with metadata and embeddings for generic knowledge base."""
    try:
        # Ensure index exists
        create_elasticsearch_index()
        
        es_client = get_elasticsearch_client()
        storage_client = get_storage_client()
        
        # Generate document ID (generic, not user-specific)
        document_id = hashlib.sha256(f"kb_{filename}_{datetime.utcnow().isoformat()}".encode()).hexdigest()
        
        # Get file content if not provided
        if not file_content:
            if not file_path:
                raise ValueError("Either file_content or file_path must be provided")
            
            with open(file_path, 'rb') as f:
                file_content = f.read()
        
        # Determine file type
        file_type = filename.split('.')[-1] if '.' in filename else 'unknown'
        
        # Create temporary file for text extraction
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Process document with layout parser (or fallback to simple extraction)
            document_data = process_document_with_layout(temp_file_path, file_type)
            
            # Extract full text content
            if isinstance(document_data, dict) and 'full_text' in document_data:
                content = document_data['full_text']
            else:
                content = document_data
            
            if not content:
                raise ValueError("Could not extract text from file")
            
            # Generate metadata from content
            metadata = generate_document_metadata(content, filename)
            
            # Generate embeddings for document chunks
            embeddings_data = generate_content_embeddings(document_data)
            
            if not embeddings_data:
                logger.warning("No embeddings generated, indexing without embeddings")
                embeddings_data = []
            
            # Prepare document for indexing with multiple chunks (generic knowledge base)
            document = {
                "document_id": document_id,
                "filename": filename,
                "file_name": filename,  # Add both for frontend compatibility
                "content": content,
                "university_name": metadata.get("university_identity", {}).get("name", ""),
                "metadata": metadata,
                "chunks": embeddings_data,  # Store all chunks with embeddings
                "num_chunks": len(embeddings_data) if embeddings_data else 0,
                "status": "indexed",
                "processing_stage": "completed",
                "indexed_at": datetime.utcnow().isoformat(),
                "upload_date": datetime.utcnow().isoformat(),  # Add upload_date for frontend
                "file_size": len(file_content),
                "file_type": file_type
            }
            
            # Index document
            try:
                logger.info(f"Attempting to index document with {len(embeddings_data)} chunks")
                logger.info(f"Sample chunk structure: {embeddings_data[0] if embeddings_data else 'No chunks'}")
                
                response = es_client.index(
                    index=ES_INDEX_NAME,
                    id=document_id,
                    body=document
                )
                
                # Check response status
                if response.get('result') not in ['created', 'updated']:
                    logger.error(f"Unexpected Elasticsearch response: {response}")
                    raise ValueError(f"Elasticsearch returned unexpected result: {response.get('result')}")
                
                logger.info(f"Document indexed successfully with {len(embeddings_data)} chunks. Response: {response}")
            except Exception as es_error:
                logger.error(f"Elasticsearch indexing error: {str(es_error)}")
                logger.error(f"Error type: {type(es_error).__name__}")
                logger.error(f"Document structure: chunks={len(embeddings_data)}, content_length={len(content)}")
                if embeddings_data:
                    logger.error(f"First chunk keys: {list(embeddings_data[0].keys())}")
                raise ValueError(f"Failed to index document in Elasticsearch: {str(es_error)}")
            
            return {
                "success": True,
                "document_id": document_id,
                "message": "Document indexed successfully",
                "metadata": {
                    "filename": filename,
                    "university_name": document['university_name'],
                    "content_length": len(content),
                    "num_chunks": len(embeddings_data) if embeddings_data else 0,
                    "index": ES_INDEX_NAME
                }
            }
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
        
    except Exception as e:
        logger.error(f"Error indexing document: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }

def search_documents(user_id, query, search_type="keyword", size=10, filters=None):
    """Search documents in generic knowledge base using Elasticsearch."""
    try:
        es_client = get_elasticsearch_client()
        
        # Build search query based on type (generic, no user_id filtering)
        if search_type == "vector":
            # Generate embedding for query (single chunk)
            query_embeddings = generate_content_embeddings(query)
            if not query_embeddings or len(query_embeddings) == 0:
                logger.warning("Could not generate query embedding, falling back to keyword search")
                search_type = "keyword"
            else:
                query_embedding = query_embeddings[0]['embedding']
            
            if not query_embedding:
                raise ValueError("Could not generate query embedding")
            
            search_body = {
                "size": size,
                "query": {
                    "bool": {
                        "must": [
                            {
                                "script_score": {
                                    "query": {"match_all": {}},  # Generic search, no user filter
                                    "script": {
                                        "source": "cosineSimilarity(params.query_vector, 'content_embedding') + 1.0",
                                        "params": {"query_vector": query_embedding}
                                    }
                                }
                            }
                        ]
                    }
                }
            }
        elif search_type == "hybrid":
            # Combine keyword and vector search
            query_embedding = generate_content_embedding(query)
            
            search_body = {
                "size": size,
                "query": {
                    "bool": {
                        "should": [
                            {
                                "multi_match": {
                                    "query": query,
                                    "fields": ["content^2", "university_name", "metadata.*"],
                                    "type": "best_fields",
                                    "fuzziness": "AUTO"
                                }
                            },
                            {
                                "script_score": {
                                    "query": {"match_all": {}},  # Generic search, no user filter
                                    "script": {
                                        "source": "cosineSimilarity(params.query_vector, 'content_embedding') + 1.0",
                                        "params": {"query_vector": query_embedding}
                                    }
                                }
                            }
                        ]
                    }
                }
            }
        else:
            # Default keyword search
            search_body = {
                "size": size,
                "query": {
                    "bool": {
                        "must": [
                            {
                                "multi_match": {
                                    "query": query,
                                    "fields": ["content^2", "university_name", "metadata.*"],
                                    "type": "best_fields",
                                    "fuzziness": "AUTO"
                                }
                            }
                        ]
                    }
                },
                "highlight": {
                    "fields": {
                        "content": {"fragment_size": 150, "number_of_fragments": 3},
                        "university_name": {}
                    }
                }
            }
        
        # Add filters if provided
        if filters:
            filter_terms = []
            for field, value in filters.items():
                filter_terms.append({"term": {field: value}})
            
            if filter_terms:
                search_body["query"]["bool"]["filter"] = filter_terms
        
        # Execute search
        response = es_client.search(index=ES_INDEX_NAME, body=search_body)
        
        # Process results
        results = []
        for hit in response['hits']['hits']:
            result = {
                "id": hit['_id'],
                "score": hit['_score'],
                "document": hit['_source'],
                "highlights": hit.get('highlight', {})
            }
            results.append(result)
        
        return {
            "success": True,
            "query": query,
            "search_type": search_type,
            "total": response['hits']['total']['value'],
            "results": results
        }
        
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }

def list_documents(user_id, size=20, from_index=0):
    """List all documents in generic knowledge base."""
    try:
        es_client = get_elasticsearch_client()
        
        search_body = {
            "size": size,
            "from": from_index,
            "query": {
                "match_all": {}  # Generic search, no user filter
            },
            "sort": [
                {"indexed_at": {"order": "desc"}}
            ]
        }
        
        response = es_client.search(index=ES_INDEX_NAME, body=search_body)
        
        documents = []
        for hit in response['hits']['hits']:
            documents.append({
                "id": hit['_id'],
                "document": hit['_source']
            })
        
        return {
            "success": True,
            "total": response['hits']['total']['value'],
            "documents": documents,
            "size": size,
            "from": from_index
        }
        
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }

def delete_document(user_id, document_id):
    """Delete a document."""
    try:
        es_client = get_elasticsearch_client()
        
        # First check if document exists and belongs to user
        try:
            doc = es_client.get(index=ES_INDEX_NAME, id=document_id)
            if doc['_source']['user_id'] != user_id:
                return {
                    "success": False,
                    "error": "Document not found or access denied"
                }
        except elasticsearch.NotFoundError:
            return {
                "success": False,
                "error": "Document not found"
            }
        
        # Delete document
        es_client.delete(index=ES_INDEX_NAME, id=document_id)
        
        return {
            "success": True,
            "message": "Document deleted successfully",
            "document_id": document_id
        }
        
    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        return {
            "success": False,
            "error": str(e)
        }

# --- Main HTTP Function (Controller) ---
@functions_framework.http
def knowledge_base_manager_es_http(request):
    """HTTP Cloud Function that acts as a controller for knowledge base operations."""
    
    # Enable CORS
    if request.method == 'OPTIONS':
        return add_cors_headers({'status': 'ok'}, 200)
    
    # Parse path to determine resource type and action
    path_parts = request.path.strip('/').split('/')
    
    if not path_parts or len(path_parts) == 0:
        return add_cors_headers({'error': 'Not Found'}, 404)
    
    resource_type = path_parts[0]
    logger.info(f"Processing {request.method} request for resource_type: {resource_type}, path: {request.path}")
    
    try:
        # --- ADMIN ROUTES ---
        if resource_type == 'admin':
            sub_path = path_parts[1] if len(path_parts) > 1 else None
            
            if sub_path == 'recreate-index' and request.method == 'POST':
                logger.info("Admin: Recreating Elasticsearch index")
                try:
                    # Delete existing index
                    delete_result = delete_elasticsearch_index()
                    # Create new index with updated schema
                    create_result = create_elasticsearch_index(force_recreate=False)
                    
                    return add_cors_headers({
                        'success': True,
                        'message': 'Index recreated successfully',
                        'deleted': delete_result,
                        'created': create_result
                    }, 200)
                except Exception as e:
                    logger.error(f"Error recreating index: {str(e)}")
                    return add_cors_headers({
                        'success': False,
                        'error': str(e)
                    }, 500)
        
        # --- DOCUMENT ROUTES ---
        if resource_type == 'documents':
            sub_path = path_parts[1:] if len(path_parts) > 1 else []
            
            # Route: /documents (POST new document)
            if len(sub_path) == 0:
                if request.method == 'POST':
                    data = request.get_json()
                    if not data or not data.get('filename'):
                        return add_cors_headers({'error': 'Filename is required'}, 400)
                    
                    user_id = data.get('user_id')
                    filename = data.get('filename')
                    
                    if not user_id:
                        return add_cors_headers({'error': 'User ID is required'}, 400)
                    
                    # Try to get file from Google Cloud Storage (generic knowledge base)
                    try:
                        storage_client = get_storage_client()
                        bucket = storage_client.bucket(BUCKET_NAME)
                        blob = bucket.blob(f"documents/{filename}")  # Generic path, no user_id
                        
                        if not blob.exists():
                            return add_cors_headers({'error': 'File not found in storage'}, 404)
                        
                        # Download file content
                        file_content = blob.download_as_bytes()
                        
                        # Index document
                        result = index_document("", filename, file_content=file_content)  # Empty user_id for generic KB
                        
                        if result['success']:
                            return add_cors_headers(result, 200)
                        else:
                            return add_cors_headers(result, 500)
                            
                    except Exception as e:
                        return add_cors_headers({'error': f'Storage error: {str(e)}'}, 500)
                
                elif request.method == 'GET':
                    # List documents (generic knowledge base)
                    data = request.args.to_dict()
                    size = int(data.get('size', 20))
                    from_index = int(data.get('from', 0))
                    
                    # Generic knowledge base - no user_id filtering
                    result = list_documents("", size, from_index)
                    
                    if result['success']:
                        return add_cors_headers(result, 200)
                    else:
                        return add_cors_headers(result, 500)
            
            # Route: /documents/search (POST search)
            elif len(sub_path) == 1 and sub_path[0] == 'search':
                if request.method == 'POST':
                    data = request.get_json()
                    if not data or not data.get('query'):
                        return add_cors_headers({'error': 'Query is required'}, 400)
                    
                    query = data.get('query')
                    search_type = data.get('search_type', 'keyword')
                    size = int(data.get('size', 10))
                    filters = data.get('filters', {})
                    
                    # Generic knowledge base - no user_id required
                    result = search_documents("", query, search_type, size, filters)
                    
                    if result['success']:
                        return add_cors_headers(result, 200)
                    else:
                        return add_cors_headers(result, 500)
            
            # Route: /documents/{document_id} (DELETE specific document)
            elif len(sub_path) == 1:
                document_id = sub_path[0]
                
                if request.method == 'DELETE':
                    data = request.get_json() if request.is_json else {}
                    user_id = data.get('user_id')
                    
                    if not user_id:
                        return add_cors_headers({'error': 'User ID is required'}, 400)
                    
                    result = delete_document(user_id, document_id)
                    
                    if result['success']:
                        return add_cors_headers(result, 200)
                    else:
                        return add_cors_headers(result, 500)
        
        # --- UPLOAD-DOCUMENT ROUTE (RAG/Firestore compatible) ---
        elif resource_type == 'upload-document':
            if request.method == 'POST':
                # Handle multipart file upload
                if 'file' not in request.files:
                    return add_cors_headers({'error': 'No file provided'}, 400)
                
                file = request.files['file']
                
                if not file.filename:
                    return add_cors_headers({'error': 'No file selected'}, 400)
                
                try:
                    # Upload file to GCS (generic knowledge base)
                    storage_client = get_storage_client()
                    bucket = storage_client.bucket(BUCKET_NAME)
                    blob = bucket.blob(f"documents/{file.filename}")  # Generic path, no user_id
                    
                    # Upload the file
                    blob.upload_from_file(file.stream, content_type=file.content_type)
                    
                    # Read file content for indexing
                    blob_content = blob.download_as_bytes()
                    
                    # Index document in Elasticsearch (generic knowledge base)
                    result = index_document("", file.filename, file_content=blob_content)  # Empty user_id for generic KB
                    
                    if result['success']:
                        return add_cors_headers(result, 200)
                    else:
                        return add_cors_headers(result, 500)
                        
                except Exception as e:
                    logger.error(f"Upload error: {str(e)}")
                    return add_cors_headers({'error': f'Upload failed: {str(e)}'}, 500)
        
        # --- DELETE ROUTE (Direct Elasticsearch deletion) ---
        elif resource_type == 'delete':
            if request.method == 'POST':
                data = request.get_json()
                logger.info(f"Delete request data: {data}")
                
                # Simple approach: get document_id from any source
                document_id = None
                
                # Try request body first
                if data:
                    document_id = data.get('document_id') or data.get('file_name') or data.get('filename') or data.get('id')
                
                # Try headers
                if not document_id:
                    document_id = request.headers.get('X-Document-ID') or request.headers.get('X-File-Name')
                
                # Try URL path
                if not document_id and len(path_parts) > 1:
                    document_id = path_parts[1]
                
                logger.info(f"Delete request - document_id: {document_id}")
                
                if not document_id:
                    return add_cors_headers({'error': 'document_id is required'}, 400)
                
                try:
                    # Direct Elasticsearch deletion
                    es_client = get_elasticsearch_client()
                    
                    # Try to delete the document (no user_id check for simplicity)
                    es_client.delete(index=ES_INDEX_NAME, id=document_id)
                    
                    logger.info(f"Successfully deleted document {document_id} from Elasticsearch")
                    
                    return add_cors_headers({
                        'success': True,
                        'message': 'Document deleted successfully',
                        'document_id': document_id
                    }, 200)
                    
                except elasticsearch.NotFoundError:
                    logger.warning(f"Document {document_id} not found in Elasticsearch")
                    return add_cors_headers({
                        'success': False,
                        'error': 'Document not found'
                    }, 404)
                    
                except Exception as e:
                    logger.error(f"Error deleting from Elasticsearch: {str(e)}")
                    return add_cors_headers({
                        'success': False,
                        'error': str(e)
                    }, 500)
        
        # --- SEARCH ROUTES ---
        elif resource_type == 'search':
            if request.method == 'POST':
                data = request.get_json()
                if not data or not data.get('query'):
                    return add_cors_headers({'error': 'Query is required'}, 400)
                
                query = data.get('query')
                search_type = data.get('search_type', 'keyword')
                size = int(data.get('size', 10))
                filters = data.get('filters', {})
                
                # Generic knowledge base - no user_id required
                result = search_documents("", query, search_type, size, filters)
                
                if result['success']:
                    return add_cors_headers(result, 200)
                else:
                    return add_cors_headers(result, 500)
        
        # --- GET DOCUMENT CONTENT ROUTE (for preview from GCS) ---
        elif resource_type == 'get-document-content' and request.method == 'POST':
            return handle_get_document_content(request)
        
        # --- HEALTH CHECK ROUTE ---
        elif resource_type == 'health':
            if request.method == 'GET':
                try:
                    # Test Elasticsearch connection
                    es_client = get_elasticsearch_client()
                    es_info = es_client.info()
                    
                    # Test Gemini connection
                    gemini_model = get_gemini_client()
                    
                    return add_cors_headers({
                        'status': 'healthy',
                        'elasticsearch': 'connected',
                        'gemini': 'connected',
                        'timestamp': datetime.utcnow().isoformat()
                    }, 200)
                except Exception as e:
                    return add_cors_headers({
                        'status': 'unhealthy',
                        'error': str(e)
                    }, 500)
        
        else:
            return add_cors_headers({'error': 'Resource not found'}, 404)
    
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return add_cors_headers({'error': 'Internal server error'}, 500)

def handle_get_document_content(request):
    """Get document content for preview from Google Cloud Storage."""
    try:
        data = request.get_json()
        logger.info(f"[GET_DOCUMENT_CONTENT] Received request data: {data}")
        
        if not data:
            logger.error(f"[GET_DOCUMENT_CONTENT ERROR] No JSON data in request")
            return add_cors_headers({
                'success': False,
                'error': 'No data provided in request'
            }, 400)
        
        file_name = data.get('file_name')
        
        if not file_name:
            logger.error(f"[GET_DOCUMENT_CONTENT ERROR] Missing file_name in data: {data}")
            return add_cors_headers({
                'success': False,
                'error': 'Missing file_name parameter'
            }, 400)
        
        logger.info(f"[GET_DOCUMENT_CONTENT] Fetching content for: {file_name}")
        
        # Get Google Cloud Storage path
        storage_path = f"knowledge-base/{file_name}"
        bucket_name = 'college-counselling-478115-knowledge-base'
        
        # Download from Google Cloud Storage
        storage_client = get_storage_client()
        if not storage_client:
            return add_cors_headers({
                'success': False,
                'error': 'Storage client not initialized'
            }, 500)
        
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(storage_path)
        
        if not blob.exists():
            logger.error(f"[GET_DOCUMENT_CONTENT ERROR] File not found: {storage_path}")
            return add_cors_headers({
                'success': False,
                'error': 'File not found in storage'
            }, 404)
        
        # Reload blob to get metadata
        blob.reload()
        
        # Get file metadata
        mime_type = blob.content_type or 'application/octet-stream'
        file_size = blob.size or 0
        upload_time = blob.time_created.strftime('%Y-%m-%d %H:%M:%S') if blob.time_created else 'Unknown'
        
        logger.info(f"[GET_DOCUMENT_CONTENT] File metadata - mime_type: {mime_type}, size: {file_size}, filename: {file_name}")
        
        # Check if it's a PDF by extension if mime_type is generic
        is_pdf = 'pdf' in mime_type.lower() or file_name.lower().endswith('.pdf')
        
        if is_pdf:
            # Make blob publicly readable temporarily and get public URL
            blob.make_public()
            download_url = blob.public_url
            content = None  # No text content for PDFs
            logger.info(f"[GET_DOCUMENT_CONTENT] Generated public URL for PDF: {download_url}")
        elif 'text' in mime_type.lower() or mime_type == 'application/json':
            try:
                content = blob.download_as_text()
                download_url = None
            except Exception as e:
                logger.error(f"[GET_DOCUMENT_CONTENT] Could not download as text: {str(e)}")
                content = f"Document: {file_name}\n\nCould not extract text content.\n\nFile size: {file_size:,} bytes"
                download_url = None
        else:
            content = f"Document: {file_name}\n\nPreview not available for this file type ({mime_type}).\n\nFile size: {file_size:,} bytes\nUploaded: {upload_time}"
            download_url = None
        
        logger.info(f"[GET_DOCUMENT_CONTENT] Successfully retrieved content (content length: {len(content) if content else 0})")
        
        return add_cors_headers({
            'success': True,
            'content': content,
            'mime_type': mime_type,
            'display_name': file_name,
            'storage_path': storage_path,
            'download_url': download_url,
            'file_size': file_size,
            'upload_time': upload_time,
            'is_pdf': is_pdf
        }, 200)
        
    except Exception as e:
        logger.error(f"[GET_DOCUMENT_CONTENT ERROR] {str(e)}")
        import traceback
        traceback.print_exc()
        return add_cors_headers({
            'success': False,
            'error': f'Failed to get content: {str(e)}'
        }, 500)

# Entry point for deployment
def knowledge_base_manager_es_http_entry(request):
    """Main entry point - redirects to the controller function."""
    return knowledge_base_manager_es_http(request)
